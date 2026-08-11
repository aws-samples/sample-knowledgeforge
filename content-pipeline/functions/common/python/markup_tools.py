"""
HTML utilities for KB article processing.

Handles extraction of plain text from HTML while preserving inline media
references (images, videos, links to sys_attachment). These references must
survive the entire pipeline — extraction, classification, quality check,
enrichment, and output — without being modified, moved, or lost.

Strategy (Option B — strip before LLM, reinsert after):
  1. Parse HTML into paragraphs, separating text from media tags
  2. Build a media map: for each media tag, record its position (paragraph index)
     and a text anchor (snippet of surrounding text) for robust reinsertion
  3. Send only clean text paragraphs to LLM — LLM never sees media tags
  4. After enrichment, reinsert media tags at the correct positions using
     paragraph index + text anchor matching
  5. Reconstruct final HTML with media tags in their original sections

Paragraph format:
  - Regular text: plain str  →  wrapped in <p> on output
  - Heading text: {"text": "...", "tag": "h2"}  →  wrapped in <h2> on output
  This lets downstream code extract text for LLM while preserving structure
  for HTML reconstruction and docx heading styles.

The media map is stored in DynamoDB alongside the article record so it's
atomic, durable, and always in sync with the article state.

Usage:
    from markup_tools import (
        extract_plain_text,
        parse_html_with_media_map,
        reinsert_media_tags,
        paragraphs_to_html,
        get_article_content,
        is_html_content,
        html_to_docx_bytes,
        get_paragraph_text,
        paragraphs_as_text_list,
    )
"""
import re
import io
import html as html_module

from bs4 import BeautifulSoup, NavigableString, Tag

try:
    from docx import Document
except ImportError:
    Document = None


# ── Constants ─────────────────────────────────────────────────────────────────

MEDIA_TAG_NAMES = {'img', 'video', 'audio', 'iframe', 'embed', 'object'}
HEADING_TAG_NAMES = {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}


# ── Placeholder functions for enrichment ──────────────────────────────────────

def extract_placeholders(html_content):
    """Replace media tags and links with [IMG_N]/[LINK_N] placeholders.
    
    Returns:
        html_with_placeholders: str — HTML with placeholders instead of media/links
        placeholder_map: list of dict
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    placeholder_map = []
    img_counter = 0
    link_counter = 0

    for tag in soup.find_all(True):
        if tag.name in MEDIA_TAG_NAMES:
            img_counter += 1
            placeholder = f'[IMG_{img_counter}]'
            para_idx = _count_preceding_blocks(tag, soup)
            placeholder_map.append({
                'placeholder': placeholder,
                'html': str(tag),
                'para_index': para_idx,
                'tag_type': tag.name,
            })
            tag.replace_with(NavigableString(placeholder))
        elif tag.name == 'a' and tag.get('href'):
            link_counter += 1
            placeholder = f'[LINK_{link_counter}]'
            para_idx = _count_preceding_blocks(tag, soup)
            placeholder_map.append({
                'placeholder': placeholder,
                'html': str(tag),
                'para_index': para_idx,
                'tag_type': 'a',
            })
            tag.replace_with(NavigableString(placeholder))

    return str(soup), placeholder_map


def _count_preceding_blocks(tag, soup):
    """Count block-level elements before this tag to estimate paragraph index."""
    block_tags = ('p', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div', 'tr')
    count = 0
    for el in soup.find_all(block_tags):
        if el.sourceline and tag.sourceline and el.sourceline >= tag.sourceline:
            break
        count += 1
    return max(0, count - 1)


def restore_placeholders(enriched_html, placeholder_map):
    """Replace [IMG_N]/[LINK_N] placeholders back with original HTML tags.
    
    Returns:
        restored_html: str — HTML with placeholders replaced by original tags
        missing: list — placeholders that were not found
    """
    result = enriched_html
    missing = []

    for entry in placeholder_map:
        placeholder = entry['placeholder']
        original_html = entry['html']
        if placeholder in result:
            result = result.replace(placeholder, original_html, 1)
        else:
            missing.append(entry)

    # Fallback: insert missing media at their para_index position
    if missing:
        soup = BeautifulSoup(result, 'html.parser')
        block_tags = ('p', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6')
        blocks = soup.find_all(block_tags)

        for entry in missing:
            para_idx = entry.get('para_index', 0)
            original_html = entry['html']
            fallback_tag = BeautifulSoup(original_html, 'html.parser')

            if para_idx < len(blocks):
                blocks[para_idx].insert_after(fallback_tag)
            else:
                soup.append(fallback_tag)

        result = str(soup)

    return result, missing

# Regex fallback for environments without BeautifulSoup
MEDIA_TAG_REGEX = re.compile(
    r'(<(?:img|video|audio|iframe|embed|object)\b[^>]*?(?:/>|>.*?</(?:img|video|audio|iframe|embed|object)>))',
    re.IGNORECASE | re.DOTALL
)
LINK_TAG_REGEX = re.compile(
    r'(<a\b[^>]*>.*?</a>)',
    re.IGNORECASE | re.DOTALL
)

# Map heading tag → docx heading level
HEADING_LEVEL_MAP = {'h1': 1, 'h2': 2, 'h3': 3, 'h4': 4, 'h5': 5, 'h6': 6}


# ── Internal helpers ──────────────────────────────────────────────────────────


def _decode_html_entities(text):
    """Decode HTML entities like &#61; back to their characters."""
    return html_module.unescape(text)


def _is_media_tag(tag):
    """Check if a BeautifulSoup tag is a media element."""
    if not Tag or not isinstance(tag, Tag):
        return False
    return tag.name in MEDIA_TAG_NAMES


def _is_attachment_link(tag):
    """Check if a BeautifulSoup tag is a link to sys_attachment or external URL."""
    if not Tag or not isinstance(tag, Tag):
        return False
    if tag.name != 'a':
        return False
    href = tag.get('href', '')
    return bool(href)


def _get_text_anchor(text, max_len=50):
    """Get a short text snippet to use as an anchor for position matching."""
    if isinstance(text, dict):
        text = text.get('text', '')
    clean = text.strip()
    if len(clean) <= max_len:
        return clean
    return clean[:max_len]


def _make_heading_para(text, tag_name):
    """Create a heading paragraph dict."""
    return {'text': text, 'tag': tag_name}


# ── Paragraph helpers (handle both str and heading dict) ──────────────────────


def get_paragraph_text(para):
    """Extract plain text from a paragraph entry (str or heading dict).
    Use this whenever you need the text content regardless of paragraph type.
    """
    if isinstance(para, dict):
        return para.get('text', '')
    return str(para)


def get_paragraph_tag(para):
    """Get the HTML tag for a paragraph. Returns 'p' for plain strings."""
    if isinstance(para, dict):
        return para.get('tag', 'p')
    return 'p'


def paragraphs_as_text_list(paragraphs, include_heading_markers=False):
    """Convert paragraphs (mixed str/dict) to a flat list of text strings.
    Use this when sending paragraphs to LLM — LLM only needs the text.
    
    If include_heading_markers=True, prefixes heading paragraphs with [H1]-[H6]
    markers so the LLM can preserve heading structure.
    """
    result = []
    for p in paragraphs:
        text = get_paragraph_text(p)
        if include_heading_markers and isinstance(p, dict) and p.get('tag') in ('h1','h2','h3','h4','h5','h6'):
            text = f"[{p['tag'].upper()}] {text}"
        result.append(text)
    return result


def _is_media_paragraph(para):
    """Check if a paragraph entry is a raw media/link HTML tag."""
    text = get_paragraph_text(para).strip()
    if not text:
        return False
    return bool(re.match(
        r'^<(?:img|video|audio|iframe|embed|object|a)\b', text, re.IGNORECASE
    ))


# ── Public API ────────────────────────────────────────────────────────────────


def extract_plain_text(html_content):
    """Strip all HTML tags and return clean plain text.
    Used for classification and embedding where we just need the words.
    Media tags are removed — use parse_html_with_media_map to preserve them.
    """
    if not html_content:
        return ''

    text = _decode_html_entities(html_content)

    if BeautifulSoup:
        soup = BeautifulSoup(text, 'html.parser')
        return soup.get_text(separator='\n', strip=True)

    # Fallback: regex-based stripping
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</?p[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</?h[1-6][^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</?li[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def parse_html_with_media_map(html_content):
    """Parse HTML content into structured paragraphs and a media map.

    Returns:
        paragraphs: list — each entry is either:
            - str: regular text paragraph (wraps in <p> on output)
            - dict {"text": "...", "tag": "h2"}: heading paragraph
        media_map: list of dict — each entry:
            {
                "tag_html": "<img src=.../>",
                "para_index": 2,
                "anchor_before": "Portal is down",
                "anchor_after": "Workaround",
                "tag_type": "img"
            }

    The media map is stored in DynamoDB and used later by
    reinsert_media_tags() to put tags back in the right place.
    """
    if not html_content:
        return [], []

    html_content = _decode_html_entities(html_content)

    if not BeautifulSoup:
        return _parse_html_regex_fallback(html_content)

    soup = BeautifulSoup(html_content, 'html.parser')

    paragraphs = []
    media_map = []
    current_text = []
    # Track the tag name of the current block element (for heading detection)
    current_block_tag = [None]
    # Track the list parent (ol/ul) for list item reconstruction
    current_list_parent = [None]

    def _flush_text(block_tag=None):
        """Flush accumulated text as a paragraph, preserving heading and list info."""
        text = ' '.join(current_text).strip()
        if text:
            tag = block_tag or current_block_tag[0]
            if tag and tag in HEADING_TAG_NAMES:
                paragraphs.append(_make_heading_para(text, tag))
            elif tag == 'li' and current_list_parent[0]:
                paragraphs.append({'text': text, 'tag': 'li', 'parent': current_list_parent[0]})
            else:
                paragraphs.append(text)
        current_text.clear()

    def _get_last_anchor():
        """Get text anchor from the last paragraph."""
        if not paragraphs:
            return ''
        return _get_text_anchor(paragraphs[-1])

    def _walk(element):
        """Walk the HTML tree, collecting text and tracking media positions."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    current_text.append(text)

            elif isinstance(child, Tag):
                # Media tag — record its position and skip its content
                if _is_media_tag(child) or _is_attachment_link(child):
                    _flush_text()
                    para_idx = len(paragraphs) - 1 if paragraphs else 0
                    media_map.append({
                        'tag_html': str(child),
                        'para_index': para_idx,
                        'anchor_before': _get_last_anchor(),
                        'anchor_after': '',
                        'tag_type': child.name,
                    })

                # Block-level elements create paragraph breaks
                elif child.name in ('p', 'div', 'li', 'tr', 'blockquote',
                                     'pre', 'section', 'article', 'ol', 'ul') or \
                     child.name in HEADING_TAG_NAMES:
                    _flush_text()
                    # Remember the block tag so nested text inherits it
                    prev_tag = current_block_tag[0]
                    prev_list = current_list_parent[0]
                    current_block_tag[0] = child.name
                    # Track ol/ul as list parent for li items
                    if child.name in ('ol', 'ul'):
                        current_list_parent[0] = child.name
                    _walk(child)
                    _flush_text(block_tag=child.name)
                    current_block_tag[0] = prev_tag
                    if child.name in ('ol', 'ul'):
                        current_list_parent[0] = prev_list

                    # Fill anchor_after for any media tags that don't have one yet
                    if paragraphs:
                        for m in media_map:
                            if not m['anchor_after']:
                                m['anchor_after'] = _get_last_anchor()

                # Inline elements — recurse without paragraph break
                else:
                    _walk(child)

    _walk(soup)
    _flush_text()

    # Fill any remaining empty anchor_after
    for m in media_map:
        if not m['anchor_after'] and paragraphs:
            m['anchor_after'] = _get_last_anchor()

    return paragraphs, media_map


def _parse_html_regex_fallback(html_content):
    """Regex-based fallback for parse_html_with_media_map when BeautifulSoup is unavailable."""
    media_map = []
    counter = 0

    def _extract_media(match):
        tag_html = match.group(0)
        tag_type = 'media'
        for name in MEDIA_TAG_NAMES:
            if f'<{name}' in tag_html.lower():
                tag_type = name
                break
        if tag_type == 'media' and '<a' in tag_html.lower():
            tag_type = 'a'
        media_map.append({
            'tag_html': tag_html,
            'para_index': counter,
            'anchor_before': '',
            'anchor_after': '',
            'tag_type': tag_type,
        })
        return ''

    # Remove media tags
    text = MEDIA_TAG_REGEX.sub(_extract_media, html_content)
    text = LINK_TAG_REGEX.sub(_extract_media, text)

    # Strip remaining HTML to get paragraphs (regex fallback loses heading info)
    plain = extract_plain_text(text)
    paragraphs = [p.strip() for p in plain.split('\n') if p.strip()]

    # Update para_index and anchors
    for i, m in enumerate(media_map):
        m['para_index'] = min(i, len(paragraphs) - 1) if paragraphs else 0
        if paragraphs:
            m['anchor_before'] = _get_text_anchor(
                paragraphs[min(m['para_index'], len(paragraphs) - 1)]
            )

    return paragraphs, media_map


def reinsert_media_tags(enriched_paragraphs, media_map, original_paragraphs=None):
    """Reinsert media tags into enriched paragraphs at their original positions.

    Uses index-based insertion: media at para_index N goes after paragraph N,
    adjusted by cumulative offset from prior insertions. This is reliable because
    our enrichment pipeline guarantees paragraph count stays the same.

    Args:
        enriched_paragraphs: list — enriched paragraphs from LLM (str or heading dict)
        media_map: list of dict — from parse_html_with_media_map / DynamoDB
        original_paragraphs: list — original paragraphs (optional, unused)

    Returns:
        list — enriched paragraphs with media tags reinserted
    """
    if not media_map:
        return enriched_paragraphs

    result = list(enriched_paragraphs)
    offset = 0  # tracks cumulative insertions to adjust para_index

    for media_entry in media_map:
        tag_html = media_entry['tag_html']
        para_idx = int(media_entry.get('para_index', 0))

        # Insert after the paragraph at para_idx, adjusted for prior insertions
        insert_pos = para_idx + offset + 1
        insert_pos = max(0, min(insert_pos, len(result)))

        result.insert(insert_pos, tag_html)
        offset += 1

    return result


def _find_position_by_anchor(paragraphs, anchor_before, anchor_after, fallback_idx):
    """Find the best paragraph index to insert a media tag."""
    if not paragraphs:
        return None

    if anchor_before:
        anchor_lower = anchor_before.lower()
        for i, para in enumerate(paragraphs):
            para_text = get_paragraph_text(para).lower()
            if anchor_lower in para_text:
                return i

    if anchor_after:
        anchor_lower = anchor_after.lower()
        for i, para in enumerate(paragraphs):
            para_text = get_paragraph_text(para).lower()
            if anchor_lower in para_text:
                return max(0, i - 1)

    return min(fallback_idx, len(paragraphs) - 1)


def paragraphs_to_html(paragraphs):
    """Convert a list of paragraphs (text + media tags + headings + list items) back to HTML.

    - Plain str paragraphs → <p>text</p>
    - Heading dicts → <h2>text</h2> (preserves original heading level)
    - List item dicts → <li>text</li> grouped in <ol>/<ul>
    - Media tag strings → passed through as-is
    """
    html_parts = []
    in_list = None  # Current open list type ('ol' or 'ul')

    for para in paragraphs:
        text = get_paragraph_text(para).strip()
        if not text:
            continue

        # Check if this is a list item
        is_li = isinstance(para, dict) and para.get('tag') == 'li'
        li_parent = para.get('parent', 'ul') if is_li else None

        # Media tags inside a list should not break the list
        is_media = _is_media_paragraph(para)

        # Close open list only if switching to non-list, non-media content
        if in_list and not is_li and not is_media:
            html_parts.append(f'</{in_list}>')
            in_list = None
        # Close and reopen if switching list type (ol -> ul or vice versa)
        if in_list and is_li and li_parent != in_list:
            html_parts.append(f'</{in_list}>')
            in_list = None

        # Media/link tag — pass through unchanged (stays inside list if open)
        if is_media:
            html_parts.append(text)
        # Heading dict — wrap in the original heading tag
        elif isinstance(para, dict) and para.get('tag') in HEADING_TAG_NAMES:
            tag = para['tag']
            html_parts.append(f'<{tag}>{text}</{tag}>')
        # List item
        elif is_li:
            if not in_list:
                html_parts.append(f'<{li_parent}>')
                in_list = li_parent
            html_parts.append(f'<li>{text}</li>')
        # Regular paragraph
        else:
            html_parts.append(f'<p>{text}</p>')

    # Close any remaining open list
    if in_list:
        html_parts.append(f'</{in_list}>')

    return '\n'.join(html_parts)


def get_article_content(article, field_preference=None):
    """Extract the content field from an article dict.
    Checks 'text' first, then 'html_content'. Returns the raw content string.

    Returns:
        (content: str, field_name: str) — the content and which field it came from
    """
    fields = field_preference or ['full_text', 'text', 'html_content']
    for field in fields:
        content = article.get(field, '')
        if content and content.strip():
            return content, field
    return '', ''


def is_html_content(content, article_type=''):
    """Determine if content is HTML based on article_type field or content inspection."""
    if article_type and article_type.upper() == 'HTML':
        return True
    if article_type and article_type.lower() == 'text':
        return False
    if content and re.search(r'<(?:p|h[1-6]|div|ul|ol|li|table|br)\b', content, re.IGNORECASE):
        return True
    return False


def html_to_docx_bytes(paragraphs, title=''):
    """Convert a list of paragraphs to a .docx file in memory.

    - Heading dicts get proper docx heading styles (Heading 1–6)
    - Plain strings become normal paragraphs
    - Media tag entries are skipped (handled separately via media map)

    Returns: bytes
    """
    if Document is None:
        raise ImportError('python-docx is required for docx conversion')

    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for para in paragraphs:
        text = get_paragraph_text(para).strip()
        if not text:
            continue
        # Skip media tags
        if _is_media_paragraph(para):
            continue
        # Heading — use proper docx heading style
        if isinstance(para, dict) and para.get('tag') in HEADING_TAG_NAMES:
            level = HEADING_LEVEL_MAP.get(para['tag'], 2)
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
